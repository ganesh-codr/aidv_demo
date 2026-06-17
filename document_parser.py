import os
import base64
import fitz  # pymupdf
import ollama
import config
from docx import Document
from openpyxl import load_workbook


def extract_text(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".docx":
        return _extract_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        return _extract_xlsx(file_path)
    else:
        return _extract_text_file(file_path)


def _ocr_page_image(pix: fitz.Pixmap) -> str:
    img_b64 = base64.b64encode(pix.tobytes("png")).decode()
    resp = ollama.chat(
        model=config.VISION_MODEL,
        messages=[{
            "role": "user",
            "content": (
                "You are an OCR assistant. Extract ALL text visible in this image exactly as "
                "written including handwritten text, printed text, table contents, labels, "
                "and captions. For tables, separate columns with | and rows with newlines. "
                "Do not interpret or summarize; only transcribe."
            ),
            "images": [img_b64],
        }]
    )
    return resp["message"]["content"].strip()


def _extract_pdf(file_path: str) -> str:
    try:
        doc = fitz.open(file_path)
        pages_text = []

        for page_num, page in enumerate(doc, start=1):
            # 1. Digital text
            text = page.get_text("text").strip()

            # 2. Tables
            table_lines = []
            try:
                tabs = page.find_tables()
                for tab in tabs.tables:
                    for row in tab.extract():
                        cells = [str(c).strip() if c else "" for c in row]
                        if any(cells):
                            table_lines.append(" | ".join(cells))
            except Exception:
                pass

            if table_lines:
                text = text + "\n\n[Tables]\n" + "\n".join(table_lines)

            # 3 & 4. OCR via Ollama vision
            has_images = bool(page.get_images())
            sparse_text = len(text.strip()) < 50

            if has_images or sparse_text:
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
                try:
                    ocr_text = _ocr_page_image(pix)
                    if sparse_text:
                        text = ocr_text
                    else:
                        text = text + "\n\n[Vision OCR - images/handwritten content]\n" + ocr_text
                except Exception as e:
                    if sparse_text:
                        msg = "[OCR skipped: vision model '{}' unavailable - {}]".format(
                            config.VISION_MODEL, e
                        )
                        text = text or msg

            if text.strip():
                pages_text.append("--- Page {} ---\n{}".format(page_num, text))
            else:
                pages_text.append("--- Page {} ---\n[No extractable content]".format(page_num))

        doc.close()
        return "\n\n".join(pages_text)

    except Exception as e:
        raise RuntimeError("Error parsing PDF file: {}".format(str(e)))


def _extract_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        text_parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = []
                seen = set()
                for cell in row.cells:
                    if cell not in seen:
                        seen.add(cell)
                        cells.append(cell.text.strip())
                if any(cells):
                    text_parts.append(" | ".join(cells))
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError("Error parsing Word file: {}".format(str(e)))


def _extract_xlsx(file_path: str) -> str:
    try:
        wb = load_workbook(file_path, data_only=True)
        text_parts = []
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append("--- Sheet: {} ---".format(sheet_name))
            for row in sheet.iter_rows(values_only=True):
                if not any(v is not None for v in row):
                    continue
                text_parts.append(" | ".join(str(v) if v is not None else "" for v in row))
        return "\n".join(text_parts)
    except Exception as e:
        raise RuntimeError("Error parsing Excel file: {}".format(str(e)))


def _extract_text_file(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise RuntimeError("Error reading text file: {}".format(str(e)))
