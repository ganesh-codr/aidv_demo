import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# PySide6 imports for PDF printing
from PySide6.QtGui import QPdfWriter, QTextDocument, QPageLayout, QPageSize
from PySide6.QtCore import QMarginsF

def markdown_to_html(md_text: str) -> str:
    """
    Convert a markdown string to HTML for QTextDocument PDF rendering.
    Supports headings, bold, italic, code blocks, lists, and tables.
    """
    lines = md_text.split("\n")
    html_parts = []
    
    in_code_block = False
    code_block_lines = []
    
    in_list = False
    list_type = None  # 'ul' or 'ol'
    
    in_table = False
    table_rows = []

    def close_list():
        nonlocal in_list, list_type
        if in_list:
            html_parts.append(f"</{list_type}>")
            in_list = False
            list_type = None

    def close_table():
        nonlocal in_table, table_rows
        if in_table:
            html_parts.append(_format_html_table(table_rows))
            table_rows = []
            in_table = False

    for line in lines:
        stripped = line.strip()
        
        # --- Code Blocks ---
        if stripped.startswith("```"):
            if in_code_block:
                # Close code block
                code_text = "\n".join(code_block_lines)
                # Escape html chars
                code_text = code_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_parts.append(f"<pre><code>{code_text}</code></pre>")
                code_block_lines = []
                in_code_block = False
            else:
                close_list()
                close_table()
                in_code_block = True
            continue
            
        if in_code_block:
            code_block_lines.append(line)
            continue

        # --- Tables ---
        # Markdown tables start/end with | or contain it
        if "|" in line:
            # Check if this is a separator line like |---|
            is_sep = re.match(r"^[\s|:-]+$", stripped)
            if is_sep:
                # We skip separator lines, but it marks that we are indeed in a table
                in_table = True
                continue
            
            close_list()
            in_table = True
            # Parse cells
            cells = [c.strip() for c in line.split("|")]
            # If the line starts and ends with |, split will create empty elements at boundaries
            if line.startswith("|"):
                cells = cells[1:]
            if line.endswith("|"):
                cells = cells[:-1]
            table_rows.append(cells)
            continue
        else:
            close_table()

        # --- Headings ---
        if stripped.startswith("#"):
            close_list()
            # Count hashes
            m = re.match(r"^(#+)\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                text = _inline_formatting(text)
                html_parts.append(f"<h{level}>{text}</h{level}>")
                continue

        # --- Lists ---
        # Bullet list
        bullet_match = re.match(r"^[\-\*\+]\s+(.*)$", stripped)
        # Ordered list
        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        
        if bullet_match:
            close_table()
            text = _inline_formatting(bullet_match.group(1))
            if not in_list or list_type != "ul":
                close_list()
                html_parts.append("<ul>")
                in_list = True
                list_type = "ul"
            html_parts.append(f"<li>{text}</li>")
            continue
        elif ordered_match:
            close_table()
            text = _inline_formatting(ordered_match.group(2))
            if not in_list or list_type != "ol":
                close_list()
                html_parts.append("<ol>")
                in_list = True
                list_type = "ol"
            html_parts.append(f"<li>{text}</li>")
            continue
        else:
            close_list()

        # --- Empty Line ---
        if not stripped:
            continue

        # --- Paragraph ---
        close_table()
        close_list()
        text = _inline_formatting(line)
        html_parts.append(f"<p>{text}</p>")

    # Clean up at end of document
    close_list()
    close_table()

    # Wrap in clean styled document
    styled_html = f"""
    <html>
    <head>
    <style>
      body {{ font-family: 'Segoe UI', Arial, sans-serif; color: #2d3748; line-height: 1.6; }}
      h1 {{ color: #1a365d; border-bottom: 2px solid #e2e8f0; padding-bottom: 6px; font-size: 20pt; margin-top: 15pt; }}
      h2 {{ color: #2b6cb0; border-bottom: 1px solid #edf2f7; padding-bottom: 4px; font-size: 16pt; margin-top: 12pt; }}
      h3 {{ color: #4a5568; font-size: 14pt; margin-top: 10pt; }}
      p {{ margin-bottom: 10pt; text-align: justify; font-size: 11pt; }}
      ul, ol {{ margin-bottom: 10pt; padding-left: 20px; font-size: 11pt; }}
      li {{ margin-bottom: 4pt; }}
      code {{ font-family: 'Courier New', monospace; background-color: #f7fafc; padding: 2px 4px; border-radius: 4px; font-size: 10pt; color: #c53030; }}
      pre {{ background-color: #f7fafc; border: 1px solid #e2e8f0; border-radius: 6px; padding: 10px; font-family: 'Courier New', monospace; font-size: 10pt; margin-bottom: 10pt; }}
      table {{ width: 100%; border-collapse: collapse; margin: 15pt 0; font-size: 10pt; }}
      th, td {{ border: 1px solid #cbd5e0; padding: 6pt 10pt; text-align: left; }}
      th {{ background-color: #edf2f7; font-weight: bold; color: #2d3748; }}
    </style>
    </head>
    <body>
      {chr(10).join(html_parts)}
    </body>
    </html>
    """
    return styled_html

def _inline_formatting(text: str) -> str:
    """Helper to replace inline markdown bold, italic, and code tags with HTML."""
    # Escape HTML tags first to prevent code blocks from parsing as tags
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    
    # Bold: **bold** or __bold__
    text = re.sub(r"\*\*(.*?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"__(.*?)__", r"<strong>\1</strong>", text)
    
    # Italic: *italic* or _italic_
    text = re.sub(r"\*(.*?)\*", r"<em>\1</em>", text)
    text = re.sub(r"_(.*?)_", r"<em>\1</em>", text)
    
    # Inline code: `code`
    text = re.sub(r"`(.*?)`", r"<code>\1</code>", text)
    
    return text

def _format_html_table(rows) -> str:
    """Format table data array into a styled HTML table."""
    if not rows:
        return ""
    
    html = ["<table>"]
    # Header
    html.append("<thead><tr>")
    for cell in rows[0]:
        html.append(f"<th>{_inline_formatting(cell)}</th>")
    html.append("</tr></thead><tbody>")
    
    # Body
    for row in rows[1:]:
        html.append("<tr>")
        # Handle cases where row might have fewer cells than header
        for i in range(len(rows[0])):
            cell_val = row[i] if i < len(row) else ""
            html.append(f"<td>{_inline_formatting(cell_val)}</td>")
        html.append("</tr>")
        
    html.append("</tbody></table>")
    return "\n".join(html)

def set_cell_background(cell, fill_hex):
    """Word utility: set background shading for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Word utility: set cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def export_to_word(markdown_content: str, output_path: str, document_title="Chatbot Answer"):
    """
    Export Markdown content to a styled Microsoft Word (.docx) file.
    Creates structured paragraphs, headings, bullet lists, code blocks, and formatted tables.
    """
    doc = Document()
    
    # Document Style Settings
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Segoe UI'
    normal_font.size = Pt(11)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run(document_title)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x36, 0x5d) # Slate Blue
    
    # Separator Line
    border_p = doc.add_paragraph()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12') # thickness
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'CBD5E0')
    pBdr.append(bottom)
    border_p._p.get_or_add_pPr().append(pBdr)

    # Parser
    lines = markdown_content.split("\n")
    
    in_code_block = False
    code_block_lines = []
    
    in_table = False
    table_rows = []

    def close_table_docx():
        nonlocal in_table, table_rows
        if in_table and table_rows:
            # Create a table
            cols_count = max(len(row) for row in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=cols_count)
            table.style = 'Table Grid'
            
            for row_idx, row in enumerate(table_rows):
                for col_idx, cell_text in enumerate(row):
                    if col_idx >= cols_count:
                        break
                    cell = table.cell(row_idx, col_idx)
                    cell.text = cell_text
                    
                    # Style cell
                    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    
                    run = p.runs[0] if p.runs else p.add_run()
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)
                    
                    if row_idx == 0:
                        # Header Row Styling
                        set_cell_background(cell, "EDF2F7") # Light gray background
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x2d, 0x37, 0x48)
            
            doc.add_paragraph() # Add empty spacing after table
            table_rows = []
            in_table = False

    for line in lines:
        stripped = line.strip()
        
        # --- Code Blocks ---
        if stripped.startswith("```"):
            if in_code_block:
                # Close code block and add to document
                code_p = doc.add_paragraph()
                code_p.paragraph_format.left_indent = Inches(0.2)
                code_p.paragraph_format.right_indent = Inches(0.2)
                code_p.paragraph_format.space_before = Pt(8)
                code_p.paragraph_format.space_after = Pt(8)
                
                # We style a custom paragraph with border/background if possible, or simple indentation and monospace font
                code_text = "\n".join(code_block_lines)
                run = code_p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xc5, 0x30, 0x30)
                
                # XML shading for paragraph background
                pPr = code_p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F7FAFC')
                pPr.append(shd)
                
                code_block_lines = []
                in_code_block = False
            else:
                close_table_docx()
                in_code_block = True
            continue
            
        if in_code_block:
            code_block_lines.append(line)
            continue

        # --- Tables ---
        if "|" in line:
            is_sep = re.match(r"^[\s|:-]+$", stripped)
            if is_sep:
                in_table = True
                continue
            
            in_table = True
            cells = [c.strip() for c in line.split("|")]
            if line.startswith("|"):
                cells = cells[1:]
            if line.endswith("|"):
                cells = cells[:-1]
            table_rows.append(cells)
            continue
        else:
            close_table_docx()

        # --- Headings ---
        if stripped.startswith("#"):
            m = re.match(r"^(#+)\s+(.*)$", stripped)
            if m:
                level = len(m.group(1))
                text = m.group(2)
                
                # Bound between 1 and 4 for heading levels
                heading_level = min(max(level, 1), 4)
                h = doc.add_heading(text, level=heading_level)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
                
                # Stylize Headings
                h_run = h.runs[0] if h.runs else h.add_run(text)
                h_run.font.name = 'Segoe UI'
                if heading_level == 1:
                    h_run.font.size = Pt(16)
                    h_run.font.color.rgb = RGBColor(0x2b, 0x6c, 0xb0)
                elif heading_level == 2:
                    h_run.font.size = Pt(14)
                    h_run.font.color.rgb = RGBColor(0x4a, 0x55, 0x68)
                else:
                    h_run.font.size = Pt(12)
                    h_run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
                continue

        # --- Bullet Lists ---
        bullet_match = re.match(r"^[\-\*\+]\s+(.*)$", stripped)
        ordered_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        
        if bullet_match:
            text = bullet_match.group(1)
            # Remove inline formatting markers
            clean_text = re.sub(r"\*\*|__|\*|_|`", "", text)
            p = doc.add_paragraph(clean_text, style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            continue
        elif ordered_match:
            text = ordered_match.group(2)
            clean_text = re.sub(r"\*\*|__|\*|_|`", "", text)
            p = doc.add_paragraph(clean_text, style='List Number')
            p.paragraph_format.space_after = Pt(3)
            continue

        # --- Empty Line ---
        if not stripped:
            continue

        # --- Standard Paragraph ---
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        
        # Simple parser for bold/inline code/normal text in a paragraph
        # Splits text by formatting tokens to format in DOCX runs
        tokens = re.split(r"(\*\*.*?\*\*|`.*?`|\*.*?\*)", line)
        for token in tokens:
            if not token:
                continue
            if token.startswith("**") and token.endswith("**"):
                r = p.add_run(token[2:-2])
                r.bold = True
            elif token.startswith("`") and token.endswith("`"):
                r = p.add_run(token[1:-1])
                r.font.name = 'Courier New'
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xc5, 0x30, 0x30)
            elif token.startswith("*") and token.endswith("*"):
                r = p.add_run(token[1:-1])
                r.italic = True
            else:
                p.add_run(token)

    close_table_docx()
    doc.save(output_path)

def export_to_pdf(markdown_content: str, output_path: str):
    """
    Render Markdown content to styled HTML and print to PDF via QPdfWriter.
    This runs inside the PyQt GUI thread.
    """
    html_content = markdown_to_html(markdown_content)
    
    writer = QPdfWriter(output_path)
    # Set high-quality resolution and paper size
    writer.setPageSize(QPageSize(QPageSize.A4))
    
    # Configure A4 margins (15mm)
    margins = QMarginsF(15, 15, 15, 15)
    layout = QPageLayout(QPageSize(QPageSize.A4), QPageLayout.Portrait, margins, QPageLayout.Millimeter)
    writer.setPageLayout(layout)
    
    doc = QTextDocument()
    doc.setHtml(html_content)
    # Set explicit document width matching printable page width (approx A4 width minus margins)
    # This ensures proper table wrap and line wrapping
    doc.setTextWidth(writer.width())
    
    # Print the document to PDF
    doc.print_(writer)

def export_to_excel(markdown_content: str, output_path: str, q_text="User Query", sheet_name="Query Response"):
    """
    Export responses to an Excel (.xlsx) sheet.
    If the markdown contains a table:
      It parses and outputs it as a structured, styled Excel grid.
    Otherwise:
      It saves the user query and the chatbot answer as rows in a Q&A worksheet.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    
    # Styles definition
    header_fill = PatternFill(start_color="1E1E2E", end_color="1E1E2E", fill_type="solid") # Dark slate
    header_font = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
    
    text_font = Font(name="Segoe UI", size=10)
    bold_font = Font(name="Segoe UI", size=10, bold=True)
    
    align_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    align_left = Alignment(horizontal="left", vertical="top", wrap_text=True)
    
    thin_border = Border(
        left=Side(style='thin', color='D2D2D2'),
        right=Side(style='thin', color='D2D2D2'),
        top=Side(style='thin', color='D2D2D2'),
        bottom=Side(style='thin', color='D2D2D2')
    )

    # Detect tables in markdown
    table_rows = []
    lines = markdown_content.split("\n")
    for line in lines:
        stripped = line.strip()
        if "|" in line:
            if re.match(r"^[\s|:-]+$", stripped):
                continue
            cells = [c.strip() for c in line.split("|")]
            if line.startswith("|"):
                cells = cells[1:]
            if line.endswith("|"):
                cells = cells[:-1]
            table_rows.append(cells)

    if table_rows:
        # Exporting markdown table grid
        for r_idx, row in enumerate(table_rows):
            for c_idx, val in enumerate(row):
                cell = ws.cell(row=r_idx+1, column=c_idx+1, value=val)
                cell.font = text_font
                cell.alignment = align_left
                cell.border = thin_border
                
                if r_idx == 0:
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.alignment = align_center
    else:
        # No table found, export Chat Log structure
        ws.column_dimensions['A'].width = 15
        ws.column_dimensions['B'].width = 80
        
        # Headers
        ws.cell(row=1, column=1, value="Role").fill = header_fill
        ws.cell(row=1, column=1).font = header_font
        ws.cell(row=1, column=1).alignment = align_center
        
        ws.cell(row=1, column=2, value="Message").fill = header_fill
        ws.cell(row=1, column=2).font = header_font
        ws.cell(row=1, column=2).alignment = align_center
        
        # User Query
        r1_a = ws.cell(row=2, column=1, value="User Query")
        r1_a.font = bold_font
        r1_a.alignment = align_left
        r1_a.border = thin_border
        
        r1_b = ws.cell(row=2, column=2, value=q_text)
        r1_b.font = text_font
        r1_b.alignment = align_left
        r1_b.border = thin_border
        
        # Chatbot Response
        r2_a = ws.cell(row=3, column=1, value="AI Response")
        r2_a.font = bold_font
        r2_a.alignment = align_left
        r2_a.border = thin_border
        
        # Clean formatting tags for Excel cells
        clean_ai_response = re.sub(r"\*\*|__|\*|_|`", "", markdown_content)
        r2_b = ws.cell(row=3, column=2, value=clean_ai_response)
        r2_b.font = text_font
        r2_b.alignment = align_left
        r2_b.border = thin_border

    # Auto-fit columns
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            val_str = str(cell.value or '')
            # Handle newlines by checking longest line in cell
            lines_in_val = val_str.split('\n')
            for l in lines_in_val:
                if len(l) > max_len:
                    max_len = len(l)
        ws.column_dimensions[col_letter].width = min(max(max_len + 4, 12), 70)

    wb.save(output_path)
